"""
Phase 1: Text Extraction Service
Extracts text from PDF, DOCX, and TXT files.
Includes OCR fallback for scanned PDFs (Tesseract).
"""

from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)


def extract_text_from_upload(file_path: str) -> str:
    """
    Extract text from uploaded document.
    Supports: PDF, DOCX, TXT
    Falls back to OCR for scanned PDFs.
    
    Args:
        file_path: Path to the uploaded file
    
    Returns:
        Extracted text content
    """
    file_ext = Path(file_path).suffix.lower()
    
    try:
        if file_ext == ".pdf":
            return extract_pdf(file_path)
        elif file_ext == ".docx":
            return extract_docx(file_path)
        elif file_ext == ".txt":
            return extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    except Exception as e:
        logger.error(f"Extraction error for {file_path}: {str(e)}")
        raise


def extract_pdf(file_path: str) -> str:
    """
    Extract text from PDF using PyMuPDF.
    Falls back to Tesseract OCR if PyMuPDF returns insufficient text.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        Extracted text
    """
    try:
        import fitz  # PyMuPDF
        
        logger.info(f"[PDF] Extracting with PyMuPDF: {file_path}")
        
        doc = fitz.open(file_path)
        text = ""
        
        for page_num, page in enumerate(doc):
            text += page.get_text()
            logger.debug(f"[PDF] Page {page_num + 1}: {len(page.get_text())} chars")
        
        doc.close()
        
        # If extraction returned minimal text, try OCR fallback
        if len(text.strip()) < 100:
            logger.warning(f"[PDF] PyMuPDF extracted only {len(text)} chars, trying OCR fallback")
            ocr_text = extract_pdf_with_ocr(file_path)
            if len(ocr_text.strip()) > len(text.strip()):
                return ocr_text
        
        return text
    
    except Exception as e:
        logger.error(f"[PDF] PyMuPDF failed: {str(e)}, trying OCR fallback")
        return extract_pdf_with_ocr(file_path)


def extract_pdf_with_ocr(file_path: str) -> str:
    """
    Extract text from PDF using Tesseract OCR.
    Used as fallback for scanned PDFs that PyMuPDF cannot read.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        OCR extracted text
    """
    try:
        import fitz  # PyMuPDF for rendering
        import pytesseract
        from PIL import Image
        import io
        
        logger.info(f"[OCR] Extracting from PDF with Tesseract: {file_path}")
        
        doc = fitz.open(file_path)
        text = ""
        
        for page_num, page in enumerate(doc):
            # Render page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # OCR the image
            page_text = pytesseract.image_to_string(img)
            text += page_text
            logger.debug(f"[OCR] Page {page_num + 1}: {len(page_text)} chars")
        
        doc.close()
        return text
    
    except ImportError:
        logger.warning("[OCR] pytesseract not installed, skipping OCR fallback")
        return ""
    except Exception as e:
        logger.error(f"[OCR] Tesseract failed: {str(e)}")
        return ""


def extract_docx(file_path: str) -> str:
    """
    Extract text from DOCX using python-docx.
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Extracted text
    """
    try:
        from docx import Document
        
        logger.info(f"[DOCX] Extracting: {file_path}")
        
        doc = Document(file_path)
        text = ""
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        logger.info(f"[DOCX] Extracted {len(text)} chars")
        return text
    
    except Exception as e:
        logger.error(f"[DOCX] Extraction failed: {str(e)}")
        raise


def extract_txt(file_path: str) -> str:
    """
    Extract text from plain TXT file.
    
    Args:
        file_path: Path to TXT file
    
    Returns:
        Extracted text
    """
    try:
        logger.info(f"[TXT] Extracting: {file_path}")
        
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        logger.info(f"[TXT] Extracted {len(text)} chars")
        return text
    
    except Exception as e:
        logger.error(f"[TXT] Extraction failed: {str(e)}")
        raise
